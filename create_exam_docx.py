#!/usr/bin/env python3
import os
import uuid
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_exam_docx(questions=None, exam_info=None, output_path=None):
    """
    Tạo file Word chứa câu hỏi trắc nghiệm như trong hình mẫu

    Parameters:
    - questions: List các câu hỏi với format [{'text': 'nội dung câu hỏi', 'options': [...], 'clo': 'CLO2', 'code': 'code mẫu'}]
    - exam_info: Dict thông tin đề thi {'title', 'subject', 'department', 'class', 'subject_code', 'time', 'date', 'exam_code'}
    - output_path: Đường dẫn lưu file output
    """
    # Tạo document mới
    doc = Document()

    # Thiết lập kích thước trang A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Sử dụng thông tin đề thi từ tham số hoặc giá trị mặc định
    if exam_info is None:
        exam_info = {
            'title': 'ĐỀ THI KẾT THÚC HỌC PHẦN\nHỌC KỲ II NĂM HỌC 2024 – 2025',
            'subject': 'Lập trình Hướng đối tượng',
            'department': 'VIỆN HỢP TÁC\nVÀ PHÁT TRIỂN ĐÀO TẠO',
            'class': '24TXTHT1 + 24TXTHC1',
            'subject_code': 'ECMP167',
            'credits': '3',
            'time': '60 phút',
            'date': '.....................',
            'exam_code': '01'
        }

    # Tạo header với bảng 2 cột
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Cột bên trái - thông tin trường và khoa
    left_cell = table.cell(0, 0)

    # Thêm logo HUTECH
    p_logo = left_cell.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run("HUTECH")
    run_logo.font.size = Pt(16)
    run_logo.font.bold = True

    # Thêm tên viện
    p_institute = left_cell.add_paragraph()
    p_institute.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_institute = p_institute.add_run(exam_info.get(
        'department', 'VIỆN HỢP TÁC\nVÀ PHÁT TRIỂN ĐÀO TẠO'))
    run_institute.font.size = Pt(12)
    run_institute.font.bold = True

    # Thêm thông tin số trang
    p_pages = left_cell.add_paragraph()
    p_pages.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Ước lượng số trang dựa trên số câu hỏi
    num_pages = len(questions) // 2 + 1 if questions else 9
    run_pages = p_pages.add_run(f"(Đề thi có {num_pages} trang)")
    run_pages.font.size = Pt(10)

    # Cột bên phải - thông tin đề thi
    right_cell = table.cell(0, 1)

    # Tên đề thi
    p_title = right_cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(exam_info.get(
        'title', 'ĐỀ THI KẾT THÚC HỌC PHẦN\nHỌC KỲ II NĂM HỌC 2024 – 2025'))
    run_title.font.size = Pt(14)
    run_title.font.bold = True

    # Thông tin đề thi
    p_info = right_cell.add_paragraph()
    p_info.add_run(
        f"Ngành/Lớp\t: {exam_info.get('class', '24TXTHT1 + 24TXTHC1')}\n")
    p_info.add_run(
        f"Tên học phần\t: {exam_info.get('subject', 'Lập trình Hướng đối tượng')}\n")
    p_info.add_run(
        f"Mã học phần\t: {exam_info.get('subject_code', 'ECMP167')}......... Số TC: {exam_info.get('credits', '3')}\n")
    p_info.add_run(
        f"Ngày thi\t: {exam_info.get('date', '.......................')}\n")
    p_info.add_run(
        f"Thời gian làm bài\t: {exam_info.get('time', '60 phút')}\n")
    p_info.add_run("Mã đề\t\t: ")
    run_code = p_info.add_run(exam_info.get('exam_code', '01'))
    run_code.font.bold = True

    # Thêm phần sử dụng tài liệu
    p_materials = right_cell.add_paragraph()
    run_materials = p_materials.add_run("SỬ DỤNG TÀI LIỆU: CÓ ☐\tKHÔNG ☒")
    run_materials.font.bold = True

    # Thêm dòng trống sau header
    doc.add_paragraph()

    # Thêm thông tin sinh viên
    p_student = doc.add_paragraph()
    p_student.add_run(
        "Họ và tên sinh viên: .....................................................................")
    p_student.add_run("Số báo danh: ..........................")

    # Thêm dòng trống
    doc.add_paragraph()

    # Nếu không có câu hỏi được cung cấp, sử dụng câu hỏi mẫu
    if not questions:
        questions = [
            {
                'text': 'Cho đoạn code sau:\nPhương thức Student (int id, String name) trong class Student gọi là:',
                'code': 'public class Student{\n    private int id;\n    private String name;\n    public Student(int id, String name){\n        this.id = id;\n        this.name = name;\n    }\n}',
                'options': [
                    'Constructor.',
                    'Destructor.',
                    'Phương thức ghi đè.',
                    'Phương thức thông thường.'
                ],
                'clo': 'CLO2'
            },
            {
                'text': 'Trong lập trình hướng đối tượng, khái niệm đóng gói (Encapsulation) có nghĩa là:',
                'options': [
                    'Ẩn thông tin và chi tiết triển khai bên trong một đối tượng.',
                    'Cho phép một lớp kế thừa thuộc tính và phương thức từ lớp khác.',
                    'Cho phép các đối tượng khác nhau phản hồi khác nhau với cùng một thông điệp.',
                    'Khả năng tạo ra các đối tượng mới từ các đối tượng có sẵn.'
                ],
                'clo': 'CLO2'
            }
        ]

    # Thêm các câu hỏi
    for i, question in enumerate(questions, 1):
        # Thêm câu hỏi
        p_q = doc.add_paragraph()
        run_q_num = p_q.add_run(f"CÂU {i}. (0.25 điểm) ")
        run_q_num.font.bold = True

        # Thêm thông tin CLO nếu có
        if question.get('clo'):
            p_q.add_run(f"({question['clo']})")

        # Thêm nội dung câu hỏi
        p_q.add_run(f"\n{question.get('text')}")

        # Nếu có đoạn code, thêm vào
        if question.get('code'):
            p_code = doc.add_paragraph()
            run_code = p_code.add_run(question['code'])
            run_code.font.name = "Courier New"
            run_code.font.color.rgb = RGBColor(0, 0, 255)

        # Thêm các đáp án
        if question.get('options'):
            labels = ["A", "B", "C", "D"]

            for j, option in enumerate(question['options']):
                if j < len(labels):
                    p_option = doc.add_paragraph()
                    p_option.paragraph_format.left_indent = Inches(0.5)
                    p_option.add_run(f"{labels[j]}. {option}")

        # Thêm dòng trống sau mỗi câu hỏi (trừ câu cuối)
        if i < len(questions):
            doc.add_paragraph()

    # Lưu file
    if output_path is None:
        output_dir = 'output'
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        output_path = os.path.join(output_dir, f'exam_{uuid.uuid4()}.docx')

    doc.save(output_path)
    print(f"Đã tạo file đề thi thành công: {output_path}")
    return output_path


def load_questions_from_json(json_file):
    """Đọc câu hỏi từ file JSON"""
    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data.get('questions', []), data.get('exam_info', {}), data.get('title', '')
    except Exception as e:
        print(f"Lỗi khi đọc file JSON: {e}")
        return [], {}, ''


if __name__ == "__main__":
    # Nếu có file JSON chứa câu hỏi được chỉ định qua tham số dòng lệnh
    import sys
    if len(sys.argv) > 1:
        questions, exam_info_json, title = load_questions_from_json(
            sys.argv[1])
        if title and not exam_info_json.get('title'):
            exam_info_json['title'] = title
        create_exam_docx(questions=questions, exam_info=exam_info_json)
    else:
        # Nếu không có tham số, sử dụng câu hỏi mẫu
        create_exam_docx()
