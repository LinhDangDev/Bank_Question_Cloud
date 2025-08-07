#!/usr/bin/env python3
"""
Exam Word Exporter using Python
Author: Linh Dang Dev

This script exports exam questions to Word document with custom formatting
"""

import sys
import json
import pyodbc
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os
from datetime import datetime
import logging

# Configure logging
logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class ExamWordExporter:
    def __init__(self, db_config):
        self.db_config = db_config
        self.connection = None

    def connect_database(self):
        """Connect to SQL Server database"""
        try:
            if self.db_config.get('trusted_connection') == 'yes':
                # Windows Authentication
                connection_string = (
                    f"DRIVER={{ODBC Driver 17 for SQL Server}};"
                    f"SERVER={self.db_config['server']};"
                    f"DATABASE={self.db_config['database']};"
                    f"Trusted_Connection=yes;"
                    "TrustServerCertificate=yes;"
                )
            else:
                # SQL Server Authentication
                connection_string = (
                    f"DRIVER={{ODBC Driver 17 for SQL Server}};"
                    f"SERVER={self.db_config['server']},{self.db_config['port']};"
                    f"DATABASE={self.db_config['database']};"
                    f"UID={self.db_config['username']};"
                    f"PWD={self.db_config['password']};"
                    "TrustServerCertificate=yes;"
                )

            self.connection = pyodbc.connect(connection_string)
            logger.info("Successfully connected to database")
            return True
        except Exception as e:
            logger.error(f"Failed to connect to database: {e}")
            return False

    def get_exam_data(self, exam_id):
        """Get exam and questions data from database"""
        try:
            cursor = self.connection.cursor()

            # Get exam info
            exam_query = """
            SELECT dt.MaDeThi, dt.TenDeThi, dt.NgayTao, dt.DaDuyet,
                   mh.TenMonHoc
            FROM DeThi dt
            LEFT JOIN MonHoc mh ON dt.MaMonHoc = mh.MaMonHoc
            WHERE dt.MaDeThi = ? AND dt.DaDuyet = 1
            """
            cursor.execute(exam_query, exam_id)
            exam_result = cursor.fetchone()

            if not exam_result:
                raise Exception(f"Approved exam not found: {exam_id}")

            exam_info = {
                'MaDeThi': exam_result[0],
                'TenDeThi': exam_result[1],
                'NgayTao': exam_result[2],
                'DaDuyet': exam_result[3],
                'TenMonHoc': exam_result[4] or 'Không có thông tin'
            }

            # Get exam questions
            questions_query = """
            SELECT ch.MaCauHoi, ch.NoiDung, ch.MaCLO, ch.CapDo,
                   ctdt.ThuTu
            FROM ChiTietDeThi ctdt
            INNER JOIN CauHoi ch ON ctdt.MaCauHoi = ch.MaCauHoi
            WHERE ctdt.MaDeThi = ?
            ORDER BY ctdt.ThuTu
            """
            cursor.execute(questions_query, exam_id)
            questions_results = cursor.fetchall()

            questions = []
            for q in questions_results:
                question = {
                    'MaCauHoi': q[0],
                    'NoiDung': q[1],
                    'MaCLO': q[2],
                    'CapDo': q[3],
                    'ThuTu': q[4],
                    'answers': []
                }

                # Get answers for this question
                answers_query = """
                SELECT MaCauTraLoi, NoiDung, LaDapAn, ThuTu
                FROM CauTraLoi
                WHERE MaCauHoi = ?
                ORDER BY ThuTu
                """
                cursor.execute(answers_query, q[0])
                answers_results = cursor.fetchall()

                for a in answers_results:
                    answer = {
                        'MaCauTraLoi': a[0],
                        'NoiDung': a[1],
                        'LaDapAn': a[2],
                        'ThuTu': a[3]
                    }
                    question['answers'].append(answer)

                questions.append(question)

            return {
                'exam': exam_info,
                'questions': questions,
                'total_questions': len(questions)
            }

        except Exception as e:
            logger.error(f"Error getting exam data: {e}")
            raise

    def create_word_document(self, exam_data, export_options):
        """Create Word document with exam content using original template"""
        try:
            # Use the original template from the old service
            template_path = os.path.join(os.path.dirname(
                __file__), '..', '..', 'template', 'TemplateHutechOffical.dotx')

            if os.path.exists(template_path):
                logger.info(f"Using original HUTECH template: {template_path}")
                # For .dotx files, we need to copy and modify
                import shutil
                temp_doc_path = template_path.replace('.dotx', '_temp.docx')
                shutil.copy2(template_path, temp_doc_path)
                doc = Document(temp_doc_path)

                # Clean up temp file after loading
                try:
                    os.remove(temp_doc_path)
                except:
                    pass
            else:
                logger.warning("Template not found, creating basic document")
                doc = Document()
                # Add basic HUTECH header if no template
                self.add_basic_hutech_header(
                    doc, exam_data['exam'], export_options)

            # Add questions
            self.add_questions(doc, exam_data['questions'], export_options)

            # Add answer key if requested
            if export_options.get('showAnswers', False):
                self.add_answer_key(
                    doc, exam_data['questions'], export_options)

            return doc

        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            raise

    def add_basic_hutech_header(self, doc, exam_info, export_options):
        """Add basic HUTECH header when template is not available"""
        # School header
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run1 = header_p.add_run("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ TP.HCM\n")
        run1.font.size = Pt(12)
        run1.bold = True

        run2 = header_p.add_run("HUTECH")
        run2.font.size = Pt(14)
        run2.bold = True

        # Exam title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(export_options.get(
            'examTitle', exam_info['TenDeThi']))
        title_run.font.size = Pt(16)
        title_run.bold = True

        # Academic year
        year_p = doc.add_paragraph()
        year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        year_run = year_p.add_run(
            f"NĂM HỌC {export_options.get('academicYear', '2024-2025')}")
        year_run.font.size = Pt(12)
        year_run.bold = True

        doc.add_paragraph()  # Space

    def add_header(self, doc, exam_info, export_options):
        """Add document header with school info and exam details"""
        # School header
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run1 = header_p.add_run("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ TP.HCM\n")
        run1.font.size = Pt(12)
        run1.bold = True

        run2 = header_p.add_run("HUTECH")
        run2.font.size = Pt(14)
        run2.bold = True

        # Exam title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(export_options.get(
            'examTitle', exam_info['TenDeThi']))
        title_run.font.size = Pt(16)
        title_run.bold = True

        # Academic year
        year_p = doc.add_paragraph()
        year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        year_run = year_p.add_run(
            f"NĂM HỌC {export_options.get('academicYear', '2024-2025')}")
        year_run.font.size = Pt(12)
        year_run.bold = True

        # Exam details table
        doc.add_paragraph()  # Space

        details_p = doc.add_paragraph()
        details_p.add_run(f"Khoa/Lớp: {export_options.get('course', '')}\t\t")
        details_p.add_run(f"Số TC: ___________\n")
        details_p.add_run(
            f"Môn thi: {export_options.get('subject', exam_info['TenMonHoc'])}\t\t")
        details_p.add_run(f"Hình thức thi: ___________\n")
        details_p.add_run(
            f"Ngày thi: {export_options.get('examDate', datetime.now().strftime('%d/%m/%Y'))}\t\t")
        details_p.add_run(f"Mã đề (Nếu có): ___________\n")
        details_p.add_run(
            f"Thời gian làm bài: {export_options.get('duration', '90 phút')}\t\t")

        # Materials checkbox
        allow_materials = export_options.get('allowMaterials', False)
        materials_text = "CÓ ☑" if allow_materials else "CÓ ☐"
        no_materials_text = "KHÔNG ☐" if allow_materials else "KHÔNG ☑"
        details_p.add_run(
            f"SỬ DỤNG TÀI LIỆU: {materials_text}  {no_materials_text}")

        doc.add_paragraph()  # Space

    def add_student_info(self, doc, export_options):
        """Add student information section"""
        student_info = export_options.get('studentInfo', {})

        info_p = doc.add_paragraph()
        info_p.add_run(
            f"Họ và tên: {student_info.get('studentName', '_' * 40)}\t\t")
        info_p.add_run(
            f"Mã số sinh viên: {student_info.get('studentId', '_' * 20)}\n")
        info_p.add_run(f"Lớp: {student_info.get('className', '_' * 30)}")

        doc.add_paragraph()  # Space

    def add_instructions(self, doc, export_options):
        """Add exam instructions"""
        instructions = export_options.get('instructions',
                                          'Thời gian làm bài: 90 phút. Không được sử dụng tài liệu.')

        inst_p = doc.add_paragraph()
        inst_run = inst_p.add_run(instructions)
        inst_run.italic = True

        doc.add_paragraph()  # Space

    def add_questions(self, doc, questions, export_options):
        """Add questions to document"""
        show_answers = export_options.get('showAnswers', False)

        for i, question in enumerate(questions, 1):
            # Question number and content
            q_p = doc.add_paragraph()
            q_num_run = q_p.add_run(f"Câu {i}: ")
            q_num_run.bold = True

            q_content_run = q_p.add_run(question['NoiDung'] or '')

            # Skip CLO and difficulty info as requested

            # Add answers
            if question['answers']:
                for j, answer in enumerate(question['answers']):
                    answer_p = doc.add_paragraph()
                    answer_p.paragraph_format.left_indent = Inches(0.3)

                    label = chr(65 + j)  # A, B, C, D
                    answer_run = answer_p.add_run(
                        f"{label}. {answer['NoiDung'] or ''}")

                    # Highlight correct answer if showing answers
                    if show_answers and answer['LaDapAn']:
                        answer_run.bold = True
                        # Add checkmark or highlight
                        correct_run = answer_p.add_run(" ✓")
                        correct_run.font.color.rgb = None  # Green color would be better
                        correct_run.bold = True

            doc.add_paragraph()  # Space between questions

    def add_answer_key(self, doc, questions, export_options):
        """Add answer key section"""
        if not export_options.get('separateAnswerSheet', False):
            return

        # Add page break
        doc.add_page_break()

        # Answer key title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run("ĐÁP ÁN")
        title_run.font.size = Pt(14)
        title_run.bold = True

        doc.add_paragraph()  # Space

        # Answer key content
        for i, question in enumerate(questions, 1):
            if question['answers']:
                correct_answers = [chr(65 + j) for j, answer in enumerate(question['answers'])
                                   if answer['LaDapAn']]
                if correct_answers:
                    answer_p = doc.add_paragraph()
                    answer_p.add_run(f"Câu {i}: {', '.join(correct_answers)}")

    def export_exam_to_word(self, exam_id, export_options, output_path):
        """Main export function"""
        try:
            logger.info(f"Starting export for exam: {exam_id}")

            # Connect to database
            if not self.connect_database():
                raise Exception("Failed to connect to database")

            # Get exam data
            exam_data = self.get_exam_data(exam_id)
            logger.info(
                f"Retrieved exam data: {exam_data['total_questions']} questions")

            # Create Word document
            doc = self.create_word_document(exam_data, export_options)

            # Save document
            doc.save(output_path)
            logger.info(f"Document saved to: {output_path}")

            return {
                'success': True,
                'message': 'Export completed successfully',
                'file_path': output_path,
                'total_questions': exam_data['total_questions']
            }

        except Exception as e:
            logger.error(f"Export failed: {e}")
            return {
                'success': False,
                'message': str(e),
                'file_path': None
            }
        finally:
            if self.connection:
                self.connection.close()


def main():
    """Main function for command line usage"""
    if len(sys.argv) < 3:
        print("Usage: python exam_word_exporter.py <exam_id> <export_options_json>")
        sys.exit(1)

    exam_id = sys.argv[1]
    export_options_json = sys.argv[2]

    try:
        export_options = json.loads(export_options_json)
    except json.JSONDecodeError as e:
        print(f"Error parsing export options JSON: {e}")
        sys.exit(1)

    # Database configuration - using Windows Authentication for local
    db_config = {
        'server': 'localhost',
        'port': 1433,
        'database': 'question_bank',
        'username': '',  # Empty for Windows Auth
        'password': '',  # Empty for Windows Auth
        'driver': 'ODBC Driver 17 for SQL Server',
        'trusted_connection': 'yes'
    }

    # Output file path
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = f"exam_export_{exam_id}_{timestamp}.docx"

    # Create exporter and run
    exporter = ExamWordExporter(db_config)
    result = exporter.export_exam_to_word(exam_id, export_options, output_path)

    # Output result as JSON
    print(json.dumps(result))


if __name__ == "__main__":
    main()
