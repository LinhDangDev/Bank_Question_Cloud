#!/usr/bin/env python3
"""
Test script để debug lỗi upload file Word
Author: Linh Dang Dev
"""

import requests
import json
import os
import sys
from pathlib import Path

# Configuration
API_BASE_URL = "http://localhost:3001/api"
TEST_FILES_DIR = Path(__file__).parent / "test_files"

def print_status(message):
    print(f"[INFO] {message}")

def print_error(message):
    print(f"[ERROR] {message}")

def print_success(message):
    print(f"[SUCCESS] {message}")

def test_api_health():
    """Test API health endpoint"""
    try:
        response = requests.get(f"{API_BASE_URL}/health", timeout=10)
        if response.status_code == 200:
            print_success("API server is running")
            return True
        else:
            print_error(f"API health check failed: {response.status_code}")
            return False
    except Exception as e:
        print_error(f"Cannot connect to API server: {e}")
        return False

def login_and_get_token():
    """Login and get authentication token"""
    try:
        login_data = {
            "username": "admin",
            "password": "123456"
        }
        
        response = requests.post(
            f"{API_BASE_URL}/auth/login",
            json=login_data,
            timeout=10
        )
        
        if response.status_code == 200:
            data = response.json()
            token = data.get('access_token')
            if token:
                print_success("Login successful")
                return token
            else:
                print_error("No token in response")
                return None
        else:
            print_error(f"Login failed: {response.status_code} - {response.text}")
            return None
            
    except Exception as e:
        print_error(f"Login error: {e}")
        return None

def test_word_upload(token, file_path):
    """Test Word file upload"""
    if not os.path.exists(file_path):
        print_error(f"Test file not found: {file_path}")
        return False
    
    try:
        headers = {
            'Authorization': f'Bearer {token}'
        }
        
        with open(file_path, 'rb') as f:
            files = {
                'file': (os.path.basename(file_path), f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            }
            
            data = {
                'processImages': 'true',
                'preserveLatex': 'true'
            }
            
            print_status(f"Uploading file: {file_path}")
            response = requests.post(
                f"{API_BASE_URL}/questions-import/upload",
                headers=headers,
                files=files,
                data=data,
                timeout=60
            )
            
            print_status(f"Response status: {response.status_code}")
            print_status(f"Response headers: {dict(response.headers)}")
            
            if response.status_code == 200:
                try:
                    result = response.json()
                    print_success("Upload successful!")
                    print(f"Questions found: {len(result.get('questions', []))}")
                    return True
                except json.JSONDecodeError:
                    print_error("Response is not valid JSON")
                    print(f"Response text: {response.text[:500]}")
                    return False
            else:
                print_error(f"Upload failed: {response.status_code}")
                print(f"Response text: {response.text}")
                return False
                
    except requests.exceptions.Timeout:
        print_error("Upload timeout - server took too long to respond")
        return False
    except Exception as e:
        print_error(f"Upload error: {e}")
        return False

def create_test_docx():
    """Create a simple test DOCX file"""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Test Questions', 0)
        
        # Add a simple question
        doc.add_paragraph('1. What is 2 + 2?')
        doc.add_paragraph('A. 3')
        doc.add_paragraph('B. 4')
        doc.add_paragraph('C. 5')
        doc.add_paragraph('D. 6')
        
        test_file = TEST_FILES_DIR / "test_simple.docx"
        TEST_FILES_DIR.mkdir(exist_ok=True)
        doc.save(test_file)
        
        print_success(f"Created test file: {test_file}")
        return str(test_file)
        
    except ImportError:
        print_error("python-docx not installed. Install with: pip install python-docx")
        return None
    except Exception as e:
        print_error(f"Error creating test file: {e}")
        return None

def main():
    print("=== Word Upload Debug Test ===")
    
    # Test API health
    if not test_api_health():
        sys.exit(1)
    
    # Login and get token
    token = login_and_get_token()
    if not token:
        sys.exit(1)
    
    # Create or find test file
    test_file = None
    
    # Try to find existing test files
    if TEST_FILES_DIR.exists():
        for file in TEST_FILES_DIR.glob("*.docx"):
            test_file = str(file)
            break
    
    # Create test file if none found
    if not test_file:
        test_file = create_test_docx()
        if not test_file:
            print_error("Cannot create test file")
            sys.exit(1)
    
    # Test upload
    print_status(f"Testing upload with file: {test_file}")
    success = test_word_upload(token, test_file)
    
    if success:
        print_success("All tests passed!")
    else:
        print_error("Upload test failed!")
        sys.exit(1)

if __name__ == "__main__":
    main()
