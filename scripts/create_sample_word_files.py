#!/usr/bin/env python3
"""
Create Sample Word Files for Question Bank Import
Author: Linh Dang Dev

This script creates sample .docx files that demonstrate the correct format
for importing questions into the Question Bank system.
"""

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("❌ python-docx library not found. Installing...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_underline(run):
    """Add underline formatting to a run"""
    run.underline = True

def create_basic_questions_sample():
    """Create a sample file with basic single-choice questions"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Sample Questions - Basic Format', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Instructions
    doc.add_paragraph("Author: Linh Dang Dev")
    doc.add_paragraph("This file demonstrates the correct format for basic single-choice questions.")
    doc.add_paragraph("Note: Correct answers are underlined.")
    doc.add_paragraph("")
    
    # Question 1
    doc.add_paragraph("Câu 1: Thủ đô của Việt Nam là gì? (CLO1)")
    p1 = doc.add_paragraph("A. ")
    run1 = p1.add_run("Hà Nội")
    add_underline(run1)  # Correct answer
    doc.add_paragraph("B. TP. Hồ Chí Minh")
    doc.add_paragraph("C. Đà Nẵng")
    doc.add_paragraph("D. Cần Thơ")
    doc.add_paragraph("")
    
    # Question 2
    doc.add_paragraph("Câu 2: Trong các ngôn ngữ lập trình sau, ngôn ngữ nào là ngôn ngữ thông dịch? (CLO2)")
    doc.add_paragraph("A. C++")
    doc.add_paragraph("B. Java")
    p2 = doc.add_paragraph("C. ")
    run2 = p2.add_run("Python")
    add_underline(run2)  # Correct answer
    doc.add_paragraph("D. C#")
    doc.add_paragraph("")
    
    # Question 3 with LaTeX
    doc.add_paragraph("Câu 3: Tính đạo hàm của hàm số f(x) = x³ + 2x² - 5x + 1 (CLO3)")
    p3 = doc.add_paragraph("A. ")
    run3 = p3.add_run("f'(x) = 3x² + 4x - 5")
    add_underline(run3)  # Correct answer
    doc.add_paragraph("B. f'(x) = 3x² + 4x + 5")
    doc.add_paragraph("C. f'(x) = x² + 4x - 5")
    doc.add_paragraph("D. f'(x) = 3x + 4")
    doc.add_paragraph("")
    
    # Question 4 with Chemistry
    doc.add_paragraph("Câu 4: Công thức phân tử của axit sulfuric là gì? (CLO1)")
    doc.add_paragraph("A. HCl")
    p4 = doc.add_paragraph("B. ")
    run4 = p4.add_run("H₂SO₄")
    add_underline(run4)  # Correct answer
    doc.add_paragraph("C. HNO₃")
    doc.add_paragraph("D. H₃PO₄")
    doc.add_paragraph("")
    
    # Question 5 - Multiple correct answers
    doc.add_paragraph("Câu 5: Những ngôn ngữ lập trình nào sau đây là ngôn ngữ hướng đối tượng? (CLO2)")
    p5a = doc.add_paragraph("A. ")
    run5a = p5a.add_run("Java")
    add_underline(run5a)  # Correct answer
    p5b = doc.add_paragraph("B. ")
    run5b = p5b.add_run("Python")
    add_underline(run5b)  # Correct answer
    doc.add_paragraph("C. C")
    p5d = doc.add_paragraph("D. ")
    run5d = p5d.add_run("C++")
    add_underline(run5d)  # Correct answer
    
    return doc

def create_group_questions_sample():
    """Create a sample file with group questions"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Sample Questions - Group Format', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Instructions
    doc.add_paragraph("Author: Linh Dang Dev")
    doc.add_paragraph("This file demonstrates the correct format for group questions.")
    doc.add_paragraph("Note: Correct answers are underlined.")
    doc.add_paragraph("")
    
    # Group Question 1
    doc.add_paragraph("Questions (1)-(4) refer to the following passage: (CLO3)")
    doc.add_paragraph("")
    doc.add_paragraph("[<sg>]")
    doc.add_paragraph("The Internet of Things (IoT) refers to the network of physical devices, vehicles, home appliances, and other items embedded with electronics, software, sensors, actuators, and connectivity which enables these objects to connect and exchange data. IoT allows objects to be sensed or controlled remotely across existing network infrastructure, creating opportunities for more direct integration of the physical world into computer-based systems.")
    doc.add_paragraph("[</sg>]")
    doc.add_paragraph("")
    
    # Child Question 1
    doc.add_paragraph("(1) What does IoT stand for?")
    doc.add_paragraph("A. Internet of Technology")
    p1 = doc.add_paragraph("B. ")
    run1 = p1.add_run("Internet of Things")
    add_underline(run1)  # Correct answer
    doc.add_paragraph("C. Integration of Technology")
    doc.add_paragraph("D. Integration of Things")
    doc.add_paragraph("")
    
    # Child Question 2
    doc.add_paragraph("(2) According to the passage, IoT enables objects to:")
    doc.add_paragraph("A. Only connect to the internet")
    p2 = doc.add_paragraph("B. ")
    run2 = p2.add_run("Be sensed or controlled remotely")
    add_underline(run2)  # Correct answer
    doc.add_paragraph("C. Replace human workers")
    doc.add_paragraph("D. Work without electricity")
    doc.add_paragraph("")
    
    # Child Question 3
    doc.add_paragraph("(3) The main benefit of IoT mentioned in the passage is:")
    doc.add_paragraph("A. Reducing costs")
    doc.add_paragraph("B. Increasing speed")
    p3 = doc.add_paragraph("C. ")
    run3 = p3.add_run("Direct integration of physical world into computer systems")
    add_underline(run3)  # Correct answer
    doc.add_paragraph("D. Eliminating human intervention")
    doc.add_paragraph("")
    
    # Child Question 4
    doc.add_paragraph("(4) Which technology contributes to IoT development?")
    p4 = doc.add_paragraph("A. ")
    run4 = p4.add_run("Wireless communication")
    add_underline(run4)  # Correct answer
    doc.add_paragraph("B. Quantum computing")
    doc.add_paragraph("C. Nuclear technology")
    doc.add_paragraph("D. Mechanical engineering")
    
    return doc

def create_fill_blank_sample():
    """Create a sample file with fill-in-blank questions"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Sample Questions - Fill-in-Blank Format', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Instructions
    doc.add_paragraph("Author: Linh Dang Dev")
    doc.add_paragraph("This file demonstrates the correct format for fill-in-blank questions.")
    doc.add_paragraph("Note: Correct answers are underlined.")
    doc.add_paragraph("")
    
    # Fill-in-blank Question 1
    doc.add_paragraph('Câu 1: Complete the sentence: "The capital of Vietnam is _____" (CLO1)')
    p1 = doc.add_paragraph("A. ")
    run1 = p1.add_run("Hanoi")
    add_underline(run1)  # Correct answer
    doc.add_paragraph("B. Ho Chi Minh City")
    doc.add_paragraph("C. Da Nang")
    doc.add_paragraph("D. Hue")
    doc.add_paragraph("")
    
    # Fill-in-blank Question 2
    doc.add_paragraph('Câu 2: Fill in the blank: "HTML stands for HyperText _____ Language" (CLO1)')
    p2 = doc.add_paragraph("A. ")
    run2 = p2.add_run("Markup")
    add_underline(run2)  # Correct answer
    doc.add_paragraph("B. Making")
    doc.add_paragraph("C. Modeling")
    doc.add_paragraph("D. Management")
    doc.add_paragraph("")
    
    # Fill-in-blank Question 3 with Math
    doc.add_paragraph('Câu 3: Complete the formula for the area of a circle: "A = _____" (CLO3)')
    doc.add_paragraph("A. πr")
    p3 = doc.add_paragraph("B. ")
    run3 = p3.add_run("πr²")
    add_underline(run3)  # Correct answer
    doc.add_paragraph("C. 2πr")
    doc.add_paragraph("D. πd")
    doc.add_paragraph("")
    
    # Fill-in-blank Question 4 with Chemistry
    doc.add_paragraph('Câu 4: Complete the chemical formula for water: "_____" (CLO1)')
    p4 = doc.add_paragraph("A. ")
    run4 = p4.add_run("H₂O")
    add_underline(run4)  # Correct answer
    doc.add_paragraph("B. CO₂")
    doc.add_paragraph("C. O₂")
    doc.add_paragraph("D. H₂O₂")
    
    return doc

def create_mixed_sample():
    """Create a sample file with mixed question types"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Sample Questions - Mixed Format', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Instructions
    doc.add_paragraph("Author: Linh Dang Dev")
    doc.add_paragraph("This file demonstrates a mix of different question types.")
    doc.add_paragraph("Note: Correct answers are underlined.")
    doc.add_paragraph("")
    
    # Single choice question
    doc.add_paragraph("Câu 1: Hệ quản trị cơ sở dữ liệu nào sau đây là mã nguồn mở? (CLO1)")
    doc.add_paragraph("A. Oracle Database")
    doc.add_paragraph("B. Microsoft SQL Server")
    p1 = doc.add_paragraph("C. ")
    run1 = p1.add_run("MySQL")
    add_underline(run1)  # Correct answer
    doc.add_paragraph("D. IBM DB2")
    doc.add_paragraph("")
    
    # Fill-in-blank question
    doc.add_paragraph('Câu 2: Fill in the blank: "The value of π (pi) is approximately _____" (CLO1)')
    p2 = doc.add_paragraph("A. ")
    run2 = p2.add_run("3.14")
    add_underline(run2)  # Correct answer
    doc.add_paragraph("B. 2.71")
    doc.add_paragraph("C. 1.62")
    doc.add_paragraph("D. 1.41")
    doc.add_paragraph("")
    
    # Group question
    doc.add_paragraph("Questions (3)-(5) refer to the following code snippet: (CLO3)")
    doc.add_paragraph("")
    doc.add_paragraph("[<sg>]")
    doc.add_paragraph("def fibonacci(n):")
    doc.add_paragraph("    if n <= 1:")
    doc.add_paragraph("        return n")
    doc.add_paragraph("    else:")
    doc.add_paragraph("        return fibonacci(n-1) + fibonacci(n-2)")
    doc.add_paragraph("[</sg>]")
    doc.add_paragraph("")
    
    # Child questions
    doc.add_paragraph("(3) What will be the output for fibonacci(3)?")
    doc.add_paragraph("A. 1")
    p3 = doc.add_paragraph("B. ")
    run3 = p3.add_run("2")
    add_underline(run3)  # Correct answer
    doc.add_paragraph("C. 3")
    doc.add_paragraph("D. 5")
    doc.add_paragraph("")
    
    doc.add_paragraph("(4) What is the base case in this recursive function?")
    doc.add_paragraph("A. n > 1")
    p4 = doc.add_paragraph("B. ")
    run4 = p4.add_run("n <= 1")
    add_underline(run4)  # Correct answer
    doc.add_paragraph("C. n == 0")
    doc.add_paragraph("D. n == 1")
    doc.add_paragraph("")
    
    doc.add_paragraph("(5) What is the time complexity of this recursive implementation?")
    doc.add_paragraph("A. O(n)")
    doc.add_paragraph("B. O(log n)")
    doc.add_paragraph("C. O(n²)")
    p5 = doc.add_paragraph("D. ")
    run5 = p5.add_run("O(2ⁿ)")
    add_underline(run5)  # Correct answer
    
    return doc

def main():
    """Main function to create all sample files"""
    print("🧪 Creating Sample Word Files for Question Bank Import")
    print("Author: Linh Dang Dev")
    print("=" * 60)
    
    # Create output directory
    output_dir = Path("sample_word_files")
    output_dir.mkdir(exist_ok=True)
    
    # Create sample files
    samples = [
        ("basic_questions_sample.docx", create_basic_questions_sample),
        ("group_questions_sample.docx", create_group_questions_sample),
        ("fill_blank_sample.docx", create_fill_blank_sample),
        ("mixed_questions_sample.docx", create_mixed_sample),
    ]
    
    for filename, create_func in samples:
        try:
            print(f"📝 Creating {filename}...")
            doc = create_func()
            filepath = output_dir / filename
            doc.save(str(filepath))
            print(f"✅ Created: {filepath}")
        except Exception as e:
            print(f"❌ Error creating {filename}: {e}")
    
    print("\n" + "=" * 60)
    print("✅ All sample files created successfully!")
    print(f"📁 Files saved in: {output_dir.absolute()}")
    print("\n📋 Usage Instructions:")
    print("1. Open the sample files in Microsoft Word")
    print("2. Review the format and underlined correct answers")
    print("3. Use these as templates for your own questions")
    print("4. Upload to Question Bank system for testing")
    print("\n⚠️  Important Notes:")
    print("- Correct answers must be underlined (not bold)")
    print("- Follow the exact format shown in samples")
    print("- Test with small files first")

if __name__ == "__main__":
    main()
